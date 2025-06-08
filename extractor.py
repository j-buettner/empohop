import os
import json
import logging
import tempfile
from typing import Dict, List, Optional, Any, Union, Tuple
import re

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class DocumentExtractor:
    """
    Class for extracting text and metadata from documents
    """
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 100):
        """
        Initialize the document extractor
        
        Args:
            chunk_size: Maximum number of tokens per chunk
            chunk_overlap: Number of overlapping tokens between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        logger.info(f"Initialized DocumentExtractor with chunk_size={chunk_size}, chunk_overlap={chunk_overlap}")
    
    def extract_and_chunk(self, source: str) -> Dict[str, Any]:
        """
        Extract text and metadata from a document and chunk it
        
        Args:
            source: Path to local file or URL
            
        Returns:
            Dictionary with document, chunks, and metadata
        """
        # Determine source type
        if source.startswith(('http://', 'https://')):
            logger.info(f"Extracting from URL: {source}")
            document = self._extract_from_url(source)
        else:
            logger.info(f"Extracting from file: {source}")
            document = self._extract_from_file(source)
        
        # Chunk the document
        chunks = self._chunk_document(document)
        
        # Extract metadata
        metadata = self._extract_metadata(document)
        
        return {
            "document": document,
            "chunks": chunks,
            "metadata": metadata
        }
    
    def _extract_from_url(self, url: str) -> Dict[str, Any]:
        """
        Extract text and metadata from a URL
        
        Args:
            url: URL to extract from
            
        Returns:
            Dictionary with text and metadata
        """
        try:
            import requests
            from bs4 import BeautifulSoup
            
            # Download the content
            response = requests.get(url)
            response.raise_for_status()
            
            # Parse HTML
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Extract text
            text = soup.get_text(separator='\n\n')
            
            # Extract metadata
            title = soup.title.string if soup.title else ""
            
            # Create document
            document = {
                "text": text,
                "metadata": {
                    "source": url,
                    "title": title,
                    "type": "web",
                }
            }
            
            return document
            
        except Exception as e:
            logger.error(f"Error extracting from URL: {str(e)}")
            raise
    
    def _extract_from_file(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text and metadata from a file
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary with text and metadata
        """
        # Get file extension
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        try:
            # Extract based on file type
            if ext == '.pdf':
                return self._extract_from_pdf(file_path)
            elif ext in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            elif ext in ['.txt', '.md', '.rst']:
                return self._extract_from_text(file_path)
            else:
                logger.warning(f"Unsupported file type: {ext}")
                return self._extract_from_text(file_path)
                
        except Exception as e:
            logger.error(f"Error extracting from file: {str(e)}")
            raise
    
    def _extract_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text and metadata from a PDF file
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Dictionary with text and metadata
        """
        try:
            # Try using PyPDF2 first
            return self._extract_from_pdf_pypdf2(file_path)
        except Exception as e:
            logger.warning(f"Error extracting with PyPDF2: {str(e)}")
            logger.info("Falling back to pdfminer.six")
            
            # Fall back to pdfminer.six
            return self._extract_from_pdf_pdfminer(file_path)
    
    def _extract_from_pdf_pypdf2(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text and metadata from a PDF file using PyPDF2
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Dictionary with text and metadata
        """
        try:
            from PyPDF2 import PdfReader
            
            # Open the PDF
            reader = PdfReader(file_path)
            
            # Extract text
            text = ""
            metadata = {}
            sections = []
            current_section = {"title": "Introduction", "content": ""}
            
            # Extract document info
            if reader.metadata:
                metadata = {
                    "title": reader.metadata.get('/Title', ''),
                    "author": reader.metadata.get('/Author', ''),
                    "subject": reader.metadata.get('/Subject', ''),
                    "creator": reader.metadata.get('/Creator', ''),
                    "producer": reader.metadata.get('/Producer', ''),
                }
            
            # Process each page
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                
                if page_text:
                    # Look for section headers
                    lines = page_text.split('\n')
                    for line in lines:
                        # Simple heuristic for section headers
                        if len(line.strip()) > 0 and len(line.strip()) < 100 and line.strip().isupper():
                            # Save the current section
                            if current_section["content"].strip():
                                sections.append(current_section)
                            
                            # Start a new section
                            current_section = {"title": line.strip(), "content": ""}
                        else:
                            current_section["content"] += line + "\n"
                    
                    text += page_text + "\n\n"
            
            # Add the last section
            if current_section["content"].strip():
                sections.append(current_section)
            
            # Create document
            document = {
                "text": text,
                "metadata": {
                    "source": file_path,
                    "title": metadata.get("title", os.path.basename(file_path)),
                    "author": metadata.get("author", ""),
                    "type": "pdf",
                    "pages": len(reader.pages),
                },
                "sections": sections
            }
            
            return document
            
        except Exception as e:
            logger.error(f"Error extracting from PDF with PyPDF2: {str(e)}")
            raise
    
    def _extract_from_pdf_pdfminer(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text and metadata from a PDF file using pdfminer.six
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Dictionary with text and metadata
        """
        try:
            from pdfminer.high_level import extract_text
            from pdfminer.pdfparser import PDFParser
            from pdfminer.pdfdocument import PDFDocument
            
            # Extract text
            text = extract_text(file_path)
            
            # Extract metadata
            with open(file_path, 'rb') as f:
                parser = PDFParser(f)
                doc = PDFDocument(parser)
                metadata = doc.info[0] if doc.info else {}
                
                # Convert metadata values from bytes to str
                metadata = {k: v.decode('utf-8', errors='ignore') if isinstance(v, bytes) else v 
                           for k, v in metadata.items()}
            
            # Extract sections
            sections = []
            current_section = {"title": "Introduction", "content": ""}
            
            # Simple section extraction based on line properties
            lines = text.split('\n')
            for line in lines:
                # Simple heuristic for section headers
                if len(line.strip()) > 0 and len(line.strip()) < 100 and line.strip().isupper():
                    # Save the current section
                    if current_section["content"].strip():
                        sections.append(current_section)
                    
                    # Start a new section
                    current_section = {"title": line.strip(), "content": ""}
                else:
                    current_section["content"] += line + "\n"
            
            # Add the last section
            if current_section["content"].strip():
                sections.append(current_section)
            
            # Create document
            document = {
                "text": text,
                "metadata": {
                    "source": file_path,
                    "title": metadata.get("Title", os.path.basename(file_path)),
                    "author": metadata.get("Author", ""),
                    "type": "pdf",
                },
                "sections": sections
            }
            
            return document
            
        except Exception as e:
            logger.error(f"Error extracting from PDF with pdfminer.six: {str(e)}")
            raise
    
    def _extract_from_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text and metadata from a DOCX file
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary with text and metadata
        """
        try:
            import docx
            
            # Open the document
            doc = docx.Document(file_path)
            
            # Extract text
            text = "\n\n".join([para.text for para in doc.paragraphs])
            
            # Extract metadata
            core_properties = doc.core_properties
            
            # Extract sections
            sections = []
            current_section = {"title": "Introduction", "content": ""}
            
            for para in doc.paragraphs:
                # Check if paragraph is a heading
                if para.style.name.startswith('Heading'):
                    # Save the current section
                    if current_section["content"].strip():
                        sections.append(current_section)
                    
                    # Start a new section
                    current_section = {"title": para.text, "content": ""}
                else:
                    current_section["content"] += para.text + "\n"
            
            # Add the last section
            if current_section["content"].strip():
                sections.append(current_section)
            
            # Create document
            document = {
                "text": text,
                "metadata": {
                    "source": file_path,
                    "title": core_properties.title or os.path.basename(file_path),
                    "author": core_properties.author or "",
                    "type": "docx",
                },
                "sections": sections
            }
            
            return document
            
        except Exception as e:
            logger.error(f"Error extracting from DOCX: {str(e)}")
            raise
    
    def _extract_from_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from a plain text file
        
        Args:
            file_path: Path to the text file
            
        Returns:
            Dictionary with text and metadata
        """
        try:
            # Read the file
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            # Extract sections
            sections = []
            current_section = {"title": "Introduction", "content": ""}
            
            lines = text.split('\n')
            for line in lines:
                # Simple heuristic for section headers
                if line.strip() and not line.strip()[0].isspace() and line.strip().endswith(':'):
                    # Save the current section
                    if current_section["content"].strip():
                        sections.append(current_section)
                    
                    # Start a new section
                    current_section = {"title": line.strip(), "content": ""}
                else:
                    current_section["content"] += line + "\n"
            
            # Add the last section
            if current_section["content"].strip():
                sections.append(current_section)
            
            # Create document
            document = {
                "text": text,
                "metadata": {
                    "source": file_path,
                    "title": os.path.basename(file_path),
                    "type": "text",
                },
                "sections": sections
            }
            
            return document
            
        except Exception as e:
            logger.error(f"Error extracting from text file: {str(e)}")
            raise
    
    def _chunk_document(self, document: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Chunk a document into smaller pieces
        
        Args:
            document: Document to chunk
            
        Returns:
            List of chunks
        """
        chunks = []
        
        # Check if document has sections
        if "sections" in document and document["sections"]:
            # Chunk by section
            for section in document["sections"]:
                section_chunks = self._chunk_text(
                    section["content"], 
                    {"section_title": section["title"]}
                )
                chunks.extend(section_chunks)
        else:
            # Chunk the entire document
            chunks = self._chunk_text(document["text"], {})
        
        # Add document metadata to each chunk
        for chunk in chunks:
            chunk["metadata"].update({
                "source": document["metadata"]["source"],
                "title": document["metadata"]["title"],
                "type": document["metadata"]["type"],
            })
        
        return chunks
    
    def _chunk_text(self, text: str, metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Chunk text into smaller pieces
        
        Args:
            text: Text to chunk
            metadata: Metadata to add to each chunk
            
        Returns:
            List of chunks
        """
        chunks = []
        
        # Split text into paragraphs
        paragraphs = text.split('\n\n')
        
        current_chunk = ""
        current_tokens = 0
        
        for paragraph in paragraphs:
            # Skip empty paragraphs
            if not paragraph.strip():
                continue
            
            # Estimate tokens (rough approximation)
            paragraph_tokens = len(paragraph.split())
            
            # If adding this paragraph would exceed the chunk size, save the current chunk
            if current_tokens + paragraph_tokens > self.chunk_size and current_chunk:
                chunks.append({
                    "text": current_chunk,
                    "metadata": metadata.copy()
                })
                
                # Start a new chunk with overlap
                overlap_tokens = min(current_tokens, self.chunk_overlap)
                current_chunk = current_chunk.split()[-overlap_tokens:]
                current_chunk = " ".join(current_chunk) + "\n\n" + paragraph
                current_tokens = overlap_tokens + paragraph_tokens
            else:
                # Add paragraph to current chunk
                if current_chunk:
                    current_chunk += "\n\n" + paragraph
                else:
                    current_chunk = paragraph
                current_tokens += paragraph_tokens
        
        # Add the last chunk
        if current_chunk:
            chunks.append({
                "text": current_chunk,
                "metadata": metadata.copy()
            })
        
        return chunks
    
    def _extract_metadata(self, document: Dict[str, Any]) -> Dict[str, Any]:
        """
        Extract metadata from a document
        
        Args:
            document: Document to extract metadata from
            
        Returns:
            Dictionary with metadata
        """
        metadata = document["metadata"].copy()
        
        # Add section information
        if "sections" in document and document["sections"]:
            metadata["sections"] = [section["title"] for section in document["sections"]]
            metadata["section_count"] = len(document["sections"])
        
        # Add text statistics
        metadata["text_length"] = len(document["text"])
        metadata["word_count"] = len(document["text"].split())
        
        return metadata
    
    def export_to_markdown(self, document: Dict[str, Any], output_path: str) -> str:
        """
        Export a document to Markdown
        
        Args:
            document: Document to export
            output_path: Path to save the Markdown file
            
        Returns:
            Path to the saved file
        """
        try:
            # Create output directory if it doesn't exist
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Generate Markdown
            markdown = f"# {document['metadata']['title']}\n\n"
            
            # Add metadata
            markdown += "## Metadata\n\n"
            for key, value in document["metadata"].items():
                if key != "title":
                    markdown += f"- **{key}**: {value}\n"
            markdown += "\n"
            
            # Add sections
            if "sections" in document and document["sections"]:
                for section in document["sections"]:
                    markdown += f"## {section['title']}\n\n"
                    markdown += section["content"] + "\n\n"
            else:
                markdown += "## Content\n\n"
                markdown += document["text"] + "\n\n"
            
            # Save to file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown)
            
            logger.info(f"Exported document to Markdown: {output_path}")
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error exporting to Markdown: {str(e)}")
            raise
    
    def save_chunks_to_json(self, chunks: List[Dict[str, Any]], output_path: str) -> str:
        """
        Save chunks to a JSON file
        
        Args:
            chunks: Chunks to save
            output_path: Path to save the JSON file
            
        Returns:
            Path to the saved file
        """
        try:
            # Create output directory if it doesn't exist
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Save to file
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(chunks, f, indent=2)
            
            logger.info(f"Saved {len(chunks)} chunks to JSON: {output_path}")
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error saving chunks to JSON: {str(e)}")
            raise
